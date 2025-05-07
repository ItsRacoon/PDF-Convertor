from flask import Flask, request, jsonify, send_from_directory, render_template, send_file
import os
import uuid
import io
from werkzeug.utils import secure_filename
from pdf2docx import Converter
import camelot
import pandas as pd
from docx import Document
from flask_cors import CORS

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF files allowed'}), 400

    format_type = request.form.get('format')
    if format_type not in ['docx', 'csv', 'xlsx']:
        return jsonify({'error': 'Invalid format'}), 400

    try:
        # Save uploaded file
        filename = secure_filename(file.filename)
        unique_id = uuid.uuid4().hex
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(upload_path)

        # Generate output filename
        output_filename = f"{os.path.splitext(filename)[0]}.{format_type}"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{unique_id}_{output_filename}")

        # Conversion
        if format_type == 'docx':
            cv = Converter(upload_path)
            # Improve conversion with better settings
            cv.convert(output_path, start=0, end=None, pages=None, 
                      kwargs={
                          'keep_image': True,  # Preserve images
                          'debug': False,
                          'multi_processing': True,  # Use multi-processing for better performance
                          'connected_border': True,  # Better handling of borders
                          'line_overlap_threshold': 0.9,  # Improve line detection
                          'line_break_width_threshold': 1.0,  # Better line break handling
                          'line_break_free_space_threshold': 0.1,  # Improve paragraph detection
                          'float_image_ignorable_gap': 5.0,  # Better image placement
                          'page_margin_factor': 0.1  # Improve margin handling
                      })
            cv.close()
        else:
            tables = camelot.read_pdf(upload_path, pages='all')
            if tables.n == 0:
                return jsonify({'error': 'No tables found in PDF'}), 400
            df = pd.concat([table.df for table in tables]) if tables.n > 1 else tables[0].df
            if format_type == 'csv':
                df.to_csv(output_path, index=False)
            else:
                df.to_excel(output_path, index=False)

        return jsonify({
            'filename': output_filename,
            'download_url': f"http://localhost:5000/download/{unique_id}_{output_filename}",
            'preview_url': f"http://localhost:5000/preview_output/{unique_id}_{output_filename}",
            'format': format_type
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['CONVERTED_FOLDER'], filename, as_attachment=True)

@app.route('/preview_output/<filename>')
def preview_output_file(filename):
    filepath = os.path.join(app.config['CONVERTED_FOLDER'], filename)
    ext = filename.rsplit('.', 1)[1].lower()
    
    if ext == 'docx':
        # Convert DOCX to HTML for preview with better formatting
        doc = Document(filepath)
        html_content = """
        <div style='font-family: Arial; padding: 20px; max-width: 800px; margin: 0 auto;'>
            <style>
                .docx-preview h1 { font-size: 24px; margin-top: 24px; margin-bottom: 16px; font-weight: bold; }
                .docx-preview h2 { font-size: 20px; margin-top: 20px; margin-bottom: 14px; font-weight: bold; }
                .docx-preview h3 { font-size: 18px; margin-top: 18px; margin-bottom: 12px; font-weight: bold; }
                .docx-preview p { margin-bottom: 12px; line-height: 1.5; }
                .docx-preview table { border-collapse: collapse; width: 100%; margin: 16px 0; }
                .docx-preview table, .docx-preview th, .docx-preview td { border: 1px solid #ddd; }
                .docx-preview th, .docx-preview td { padding: 8px; text-align: left; }
                .docx-preview th { background-color: #f2f2f2; }
                .docx-preview img { max-width: 100%; height: auto; }
                .docx-preview .bold { font-weight: bold; }
                .docx-preview .italic { font-style: italic; }
                .docx-preview .underline { text-decoration: underline; }
            </style>
            <div class="docx-preview">
        """
        
        # Process paragraphs with better style handling
        for para in doc.paragraphs:
            if not para.text.strip():  # Skip empty paragraphs
                html_content += "<p>&nbsp;</p>"
                continue
                
            # Determine paragraph style
            if para.style.name.startswith('Heading 1'):
                html_content += f"<h1>{para.text}</h1>"
            elif para.style.name.startswith('Heading 2'):
                html_content += f"<h2>{para.text}</h2>"
            elif para.style.name.startswith('Heading 3'):
                html_content += f"<h3>{para.text}</h3>"
            else:
                # Process runs to maintain formatting
                if para.runs:
                    p_html = "<p>"
                    for run in para.runs:
                        run_text = run.text
                        if run.bold:
                            run_text = f'<span class="bold">{run_text}</span>'
                        if run.italic:
                            run_text = f'<span class="italic">{run_text}</span>'
                        if run.underline:
                            run_text = f'<span class="underline">{run_text}</span>'
                        p_html += run_text
                    p_html += "</p>"
                    html_content += p_html
                else:
                    html_content += f"<p>{para.text}</p>"
        
        # Process tables
        for table in doc.tables:
            table_html = "<table>"
            for i, row in enumerate(table.rows):
                table_html += "<tr>"
                for cell in row.cells:
                    tag = "th" if i == 0 else "td"
                    table_html += f"<{tag}>{cell.text}</{tag}>"
                table_html += "</tr>"
            table_html += "</table>"
            html_content += table_html
            
        html_content += "</div></div>"
        return html_content
    elif ext == 'csv':
        # Read CSV and convert to HTML table
        df = pd.read_csv(filepath)
        return df.to_html(classes='data-table', index=False)
    elif ext == 'xlsx':
        # Read Excel and convert to HTML table
        df = pd.read_excel(filepath)
        return df.to_html(classes='data-table', index=False)
    else:
        return "Preview not available for this file type"

if __name__ == '__main__':
    app.run(debug=True)