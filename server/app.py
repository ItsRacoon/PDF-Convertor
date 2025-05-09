from flask import Flask, request, jsonify, send_from_directory
import os
import uuid
import logging
import sys
from werkzeug.utils import secure_filename
from pdf2docx import Converter
import pandas as pd
from docx import Document
from flask_cors import CORS
from pathlib import Path

# Try to import camelot, but don't fail if it's not available
try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    logging.warning("Camelot not available. Will use tabula-py for table extraction.")

# Load environment variables from .env file if present
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

# Check if running on Render.com
IS_RENDER = os.environ.get('RENDER_ENV') == 'production'
USE_TABULA_FALLBACK = os.environ.get('TABULA_FALLBACK', 'false').lower() == 'true'

logger.info(f"Running in {'Render.com' if IS_RENDER else 'local'} environment")
logger.info(f"Tabula fallback enabled: {USE_TABULA_FALLBACK}")

# Initialize Flask app
app = Flask(__name__)
CORS(app)

app.config.update(
    UPLOAD_FOLDER=os.path.join(os.getcwd(), 'uploads'),
    CONVERTED_FOLDER=os.path.join(os.getcwd(), 'converted'),
    DEBUG=os.environ.get('FLASK_ENV') == 'development',
    MAX_CONTENT_LENGTH=16 * 1024 * 1024,
    ALLOWED_EXTENSIONS={'pdf'},
    CONVERSION_FORMATS={'docx', 'csv', 'xlsx'}
)

# Configure basic logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create directories
Path(app.config['UPLOAD_FOLDER']).mkdir(exist_ok=True)
Path(app.config['CONVERTED_FOLDER']).mkdir(exist_ok=True)

def allowed_file(filename):
    """Check if uploaded file has allowed extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/')
def home():
    """API root endpoint"""
    return jsonify({
        'name': 'PDF Convertor API',
        'version': '1.0.0',
        'description': 'API for converting PDF files to various formats',
        'endpoints': {
            '/convert': 'POST - Convert PDF to selected format',
            '/download/<filename>': 'GET - Download converted file',
            '/preview_output/<filename>': 'GET - Preview converted file',
            '/health': 'GET - API health check'
        }
    })

@app.route('/convert', methods=['POST'])
def convert_file():
    """Handle PDF conversion to selected format"""
    # Validate request
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid or missing file'}), 400

    format_type = request.form.get('format')
    if format_type not in app.config['CONVERSION_FORMATS']:
        return jsonify({'error': f"Invalid format"}), 400

    try:
        # Setup filenames
        filename = secure_filename(file.filename)
        unique_id = uuid.uuid4().hex
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        output_filename = f"{os.path.splitext(filename)[0]}.{format_type}"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{unique_id}_{output_filename}")
        
        # Save uploaded file
        file.save(upload_path)
        
        # Convert based on format
        if format_type == 'docx':
            cv = Converter(upload_path)
            cv.convert(output_path, start=0, end=None, 
                      kwargs={
                          'keep_image': True,
                          'multi_processing': True
                      })
            cv.close()
        else:  # csv or xlsx
            # Determine which extraction method to try first
            if IS_RENDER or not CAMELOT_AVAILABLE or USE_TABULA_FALLBACK:
                # On Render, or if camelot is not available, try tabula-py first
                extraction_methods = ['tabula', 'camelot'] if CAMELOT_AVAILABLE else ['tabula']
            else:
                # Locally, try camelot first (if Ghostscript is installed)
                extraction_methods = ['camelot', 'tabula']
                
            logger.info(f"Using extraction methods in order: {extraction_methods}")
            
            last_error = None
            for method in extraction_methods:
                try:
                    if method == 'camelot' and CAMELOT_AVAILABLE:
                        # Try using camelot (requires Ghostscript)
                        logger.info("Attempting table extraction with camelot...")
                        tables = camelot.read_pdf(upload_path, pages='all')
                        if tables.n == 0:
                            logger.warning("No tables found with camelot")
                            continue
                            
                        df = pd.concat([table.df for table in tables]) if tables.n > 1 else tables[0].df
                    
                    elif method == 'tabula':
                        # Try using tabula-py
                        logger.info("Attempting table extraction with tabula-py...")
                        import importlib
                        tabula_spec = importlib.util.find_spec('tabula')
                        
                        if tabula_spec is None:
                            # Install tabula-py if not available
                            import subprocess
                            logger.info("Installing tabula-py...")
                            subprocess.check_call(['pip', 'install', 'tabula-py'])
                        
                        import tabula
                        dfs = tabula.read_pdf(upload_path, pages='all')
                        if not dfs:
                            logger.warning("No tables found with tabula-py")
                            continue
                            
                        df = pd.concat(dfs) if len(dfs) > 1 else dfs[0]
                    
                    # If we got here, extraction was successful
                    if format_type == 'csv':
                        df.to_csv(output_path, index=False)
                    else:  # xlsx
                        df.to_excel(output_path, index=False)
                    
                    # Successfully extracted and saved
                    logger.info(f"Successfully extracted tables using {method}")
                    break
                    
                except Exception as e:
                    last_error = e
                    error_msg = str(e)
                    logger.warning(f"{method} extraction failed: {error_msg}")
                    continue
            
            # If we've tried all methods and still failed, try to extract text as a fallback
            if 'df' not in locals():
                logger.warning("All table extraction methods failed, attempting text extraction fallback")
                try:
                    # Try to extract text from PDF using pdfplumber or PyMuPDF
                    try:
                        import pdfplumber
                        
                        logger.info("Extracting text with pdfplumber as fallback")
                        with pdfplumber.open(upload_path) as pdf:
                            text_content = []
                            for i, page in enumerate(pdf.pages):
                                page_text = page.extract_text() or ""
                                if page_text.strip():
                                    text_content.append([f"Page {i+1}", page_text.strip()])
                            
                            if text_content:
                                # Create a DataFrame with the extracted text
                                df = pd.DataFrame(text_content, columns=["Page", "Content"])
                                logger.info("Created text-based table as fallback")
                            else:
                                # If no text was extracted, try PyMuPDF as a second fallback
                                raise ValueError("No text extracted with pdfplumber")
                    except Exception as pdfplumber_error:
                        # Try PyMuPDF as a fallback
                        logger.info(f"Pdfplumber failed: {str(pdfplumber_error)}. Trying PyMuPDF...")
                        try:
                            import fitz  # PyMuPDF
                            
                            text_content = []
                            doc = fitz.open(upload_path)
                            for i in range(len(doc)):
                                page = doc[i]
                                page_text = page.get_text()
                                if page_text.strip():
                                    text_content.append([f"Page {i+1}", page_text.strip()])
                            
                            if text_content:
                                # Create a DataFrame with the extracted text
                                df = pd.DataFrame(text_content, columns=["Page", "Content"])
                                logger.info("Created text-based table using PyMuPDF fallback")
                            else:
                                # If no text was extracted, create an empty DataFrame with a message
                                df = pd.DataFrame([["No extractable content found in this PDF"]], 
                                                columns=["Message"])
                                logger.warning("No text content found in PDF with PyMuPDF either")
                        except Exception as pymupdf_error:
                            logger.error(f"PyMuPDF fallback also failed: {str(pymupdf_error)}")
                            # Create a simple DataFrame with a message
                            df = pd.DataFrame([["This PDF does not contain extractable text or tables"]], 
                                            columns=["Message"])
                            logger.warning("Created empty DataFrame as last resort")
                    
                    # Save the fallback content
                    if format_type == 'csv':
                        df.to_csv(output_path, index=False)
                    else:  # xlsx
                        df.to_excel(output_path, index=False)
                        
                except Exception as e:
                    # If even the text extraction fallback fails, return an error
                    logger.error(f"Text extraction fallback also failed: {str(e)}")
                    
                    # Check for common error conditions
                    error_message = 'PDF extraction failed.'
                    if 'password' in str(e).lower() or 'decrypt' in str(e).lower():
                        error_message += ' The PDF appears to be password-protected.'
                    elif 'corrupt' in str(e).lower() or 'invalid' in str(e).lower():
                        error_message += ' The PDF file may be corrupted.'
                    else:
                        error_message += ' The PDF may be in an unsupported format or have restricted permissions.'
                    
                    return jsonify({
                        'error': error_message
                    }), 500
        
        # Generate response
        host_url = request.host_url.rstrip('/')
        return jsonify({
            'filename': output_filename,
            'download_url': f"{host_url}/download/{unique_id}_{output_filename}",
            'preview_url': f"{host_url}/preview_output/{unique_id}_{output_filename}",
            'format': format_type
        })

    except Exception as e:
        logger.error(f"Conversion error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Serve file for download"""
    return send_from_directory(app.config['CONVERTED_FOLDER'], filename, as_attachment=True)

@app.route('/preview_output/<filename>')
def preview_output_file(filename):
    """Generate HTML preview of converted files"""
    filepath = os.path.join(app.config['CONVERTED_FOLDER'], filename)
    if not os.path.exists(filepath):
        return "File not found", 404
        
    ext = filename.rsplit('.', 1)[1].lower()
    
    try:
        if ext == 'docx':
            # Generate DOCX preview
            doc = Document(filepath)
            html_content = """
            <style>
                .docx-preview { font-family: Arial; padding: 20px; max-width: 800px; margin: 0 auto; }
                .docx-preview h1 { font-size: 24px; font-weight: bold; }
                .docx-preview h2 { font-size: 20px; font-weight: bold; }
                .docx-preview p { margin-bottom: 12px; line-height: 1.5; }
                .docx-preview table { border-collapse: collapse; width: 100%; margin: 16px 0; }
                .docx-preview table, .docx-preview th, .docx-preview td { border: 1px solid #ddd; padding: 8px; }
                .bold { font-weight: bold; }
                .italic { font-style: italic; }
                .underline { text-decoration: underline; }
            </style>
            <div class="docx-preview">
            """
            
            # Process paragraphs
            for para in doc.paragraphs:
                if not para.text.strip():
                    html_content += "<p>&nbsp;</p>"
                    continue
                    
                if para.style.name.startswith('Heading 1'):
                    html_content += f"<h1>{para.text}</h1>"
                elif para.style.name.startswith('Heading 2'):
                    html_content += f"<h2>{para.text}</h2>"
                else:
                    # Process runs with formatting
                    if para.runs:
                        p_html = "<p>"
                        for run in para.runs:
                            run_text = run.text.replace('<', '&lt;').replace('>', '&gt;')
                            if run.bold: run_text = f'<span class="bold">{run_text}</span>'
                            if run.italic: run_text = f'<span class="italic">{run_text}</span>'
                            if run.underline: run_text = f'<span class="underline">{run_text}</span>'
                            p_html += run_text
                        p_html += "</p>"
                        html_content += p_html
                    else:
                        html_content += f"<p>{para.text}</p>"
            
            # Process tables
            for table in doc.tables:
                table_html = "<table>"
                for i, row in enumerate(table.rows):
                    table_html += "<tr>"
                    for cell in row.cells:
                        tag = "th" if i == 0 else "td"
                        table_html += f"<{tag}>{cell.text}</{tag}>"
                    table_html += "</tr>"
                table_html += "</table>"
                html_content += table_html
                
            html_content += "</div>"
            return html_content
            
        elif ext in ['csv', 'xlsx']:
            # Generate table preview
            if ext == 'csv':
                df = pd.read_csv(filepath)
            else:  # xlsx
                df = pd.read_excel(filepath)
                
            # Basic styled table
            html = """
            <style>
                .data-table { border-collapse: collapse; width: 100%; }
                .data-table th { background-color: #f2f2f2; font-weight: bold; text-align: left; }
                .data-table th, .data-table td { border: 1px solid #ddd; padding: 8px; }
                .data-table tr:nth-child(even) { background-color: #f9f9f9; }
            </style>
            """
            html += df.to_html(classes='data-table', index=False, na_rep='')
            return html
            
        else:
            return "Preview not available for this file type", 400
            
    except Exception as e:
        logger.error(f"Preview error: {str(e)}")
        return f"Error generating preview: {str(e)}", 500

@app.route('/health')
def health_check():
    """API health check endpoint"""
    return jsonify({'status': 'healthy'}), 200

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))